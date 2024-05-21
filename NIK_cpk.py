import numpy as np
import pandas as pd
from docx import Document
from datetime import datetime

print('---cpk數值生成---')



def generate_values(USL, LSL, cpk, n):
    C = (USL + LSL) / 2
    T = USL - LSL
    sigma = T / (6 * cpk)
    
    # 生成符合Cpk的數據
    while True:
        values = np.random.normal(loc=C, scale=sigma, size=n)
        X_bar = np.mean(values)
        calculated_sigma = np.std(values, ddof=1)
        Cp = (USL - LSL) / (6 * calculated_sigma)
        Ca = abs((X_bar - C) / (T / 2))
        calculated_cpk = Cp * (1 - Ca)
        
        if np.isclose(calculated_cpk, cpk, atol=0.01):
            break
            
    return values

while True:
    # 輸入參數
    decimal_places = int(input("請問要保留小數點幾位數："))
    USL = float(input("請輸入規格上限 (USL): "))
    LSL = float(input("請輸入規格下限 (LSL): "))
    cpk = float(input("請輸入Cpk值: "))
    n = int(input("請輸入想要生成的數據筆數: "))

    # 生成數據
    values = generate_values(USL, LSL, cpk, n)
    print("生成的數值為:")
    for value in values:
        print(f"{value:.{decimal_places}f}")

    # 获取当前时间并格式化为字符串
    current_time = datetime.now().strftime("%Y.%m.%d.%H.%M.%S")

    # 询问是否需要输出格式
    format_choice = input("请选择输出格式：1. 文本文件 (.txt)  2. Excel 表格 (.xlsx)  3. Word 文档 (.docx)  4. 不需要输出格式\n")

    # 输出到文本文件
    if format_choice == '1':
        with open(f'output_{current_time}.txt', 'w', encoding='utf-8') as f:
            f.write("生成的數值为:\n")
            for value in values:
                f.write(f"{value:.{decimal_places}f}\n")
        print(f"数据已输出到 output_{current_time}.txt")

    # 输出到Excel表格
    elif format_choice == '2':
        df = pd.DataFrame(values, columns=['Values'])
        df.to_excel(f'output_{current_time}.xlsx', index=False)
        print(f"数据已输出到 output_{current_time}.xlsx")

    # 输出到Word文档
    elif format_choice == '3':
        doc = Document()
        doc.add_heading('Generated Values', level=1)
        for value in values:
            doc.add_paragraph(f"{value:.{decimal_places}f}")
        doc.save(f'output_{current_time}.docx')
        print(f"数据已输出到 output_{current_time}.docx")

    # 不需要输出格式
    elif format_choice == '4':
        print("不需要输出格式")
    else:
        print("錯誤輸入, 默認不需要输出格式")
    
    # 询问是否需要继续生成数据
    choice = input("是否需要继续生成数据？(输入 'Y' 继续，输入其他内容退出): ")
    if choice.lower() != 'y':
        break
